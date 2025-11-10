"""Document processing utilities for extracting content from various file formats."""
import os
from typing import List
from pathlib import Path
from models import SupportingDocument


class DocumentProcessor:
    """Process and extract content from supporting documents."""
    
    def __init__(self):
        self.supported_formats = ['.txt', '.md', '.pdf', '.docx']
    
    def process_document(self, file_path: str) -> SupportingDocument:
        """
        Process a single document and extract its content.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            SupportingDocument object with extracted content
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        extension = file_path.suffix.lower()
        
        if extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {extension}")
        
        # Extract content based on file type
        if extension in ['.txt', '.md']:
            content = self._process_text_file(file_path)
        elif extension == '.pdf':
            content = self._process_pdf(file_path)
        elif extension == '.docx':
            content = self._process_docx(file_path)
        else:
            content = ""
        
        return SupportingDocument(
            filename=file_path.name,
            content=content,
            document_type=self._get_document_type(file_path)
        )
    
    def process_multiple_documents(self, file_paths: List[str]) -> List[SupportingDocument]:
        """
        Process multiple documents.
        
        Args:
            file_paths: List of paths to document files
            
        Returns:
            List of SupportingDocument objects
        """
        documents = []
        for file_path in file_paths:
            try:
                doc = self.process_document(file_path)
                documents.append(doc)
            except Exception as e:
                print(f"Error processing {file_path}: {e}")
        
        return documents
    
    def _process_text_file(self, file_path: Path) -> str:
        """Extract content from text/markdown files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            return f"Error reading text file: {e}"
    
    def _process_pdf(self, file_path: Path) -> str:
        """Extract content from PDF files."""
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            content = []
            for page in reader.pages:
                content.append(page.extract_text())
            return "\n\n".join(content)
        except ImportError:
            return "PDF processing requires 'pypdf' package. Install with: pip install pypdf"
        except Exception as e:
            return f"Error reading PDF file: {e}"
    
    def _process_docx(self, file_path: Path) -> str:
        """Extract content from DOCX files."""
        try:
            from docx import Document
            doc = Document(file_path)
            content = []
            for para in doc.paragraphs:
                content.append(para.text)
            return "\n\n".join(content)
        except ImportError:
            return "DOCX processing requires 'python-docx' package. Install with: pip install python-docx"
        except Exception as e:
            return f"Error reading DOCX file: {e}"
    
    def _get_document_type(self, file_path: Path) -> str:
        """Determine the document type based on filename and extension."""
        filename = file_path.name.lower()
        
        if 'requirement' in filename or 'req' in filename:
            return 'requirements'
        elif 'spec' in filename:
            return 'specification'
        elif 'design' in filename:
            return 'design'
        elif 'api' in filename:
            return 'api_documentation'
        else:
            return 'general'
